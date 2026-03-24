import os
import logging
import nltk
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)
try:
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    nltk.download('punkt_tab', quiet=True)

import fitz  # pymupdf
from docx import Document
from nltk.tokenize import sent_tokenize

logger = logging.getLogger(__name__)

def stream_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        while True:
            block = f.read(1024 * 1024)
            if not block:
                break
            yield block, None

def stream_pdf(path):
    doc = fitz.open(path)
    for page_num, page in enumerate(doc):
        text = page.get_text()
        if text and text.strip():
            yield text, page_num
    doc.close()

def stream_docx(path):
    doc = Document(path)
    for p in doc.paragraphs:
        if p.text.strip():
            yield p.text, None

def _table_to_markdown(rows):
    """Convert a list of rows (each a list of cell strings) to markdown table."""
    if not rows or len(rows) < 1:
        return ""
    # Clean cells: replace None with empty string, strip whitespace
    clean = []
    for row in rows:
        clean.append([str(cell).strip() if cell else "" for cell in row])

    # Build markdown: first row as header
    header = clean[0]
    md_lines = ["| " + " | ".join(header) + " |"]
    md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
    for row in clean[1:]:
        # Pad row if shorter than header
        padded = row + [""] * (len(header) - len(row))
        md_lines.append("| " + " | ".join(padded[:len(header)]) + " |")
    return "\n".join(md_lines)


def extract_pdf_tables(path):
    """Extract tables from PDF pages using PyMuPDF's find_tables().
    Yields (markdown_text, page_num, table_index) for each table found."""
    doc = fitz.open(path)
    for page_num, page in enumerate(doc):
        try:
            tables = page.find_tables()
        except AttributeError:
            # PyMuPDF version too old for find_tables()
            break
        for table_idx, table in enumerate(tables):
            try:
                rows = table.extract()
                md = _table_to_markdown(rows)
                if md.strip():
                    yield md, page_num, table_idx
            except Exception:
                continue
    doc.close()


def extract_docx_tables(path):
    """Extract tables from DOCX files using python-docx.
    Yields (markdown_text, None, table_index) for each table found."""
    doc = Document(path)
    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        md = _table_to_markdown(rows)
        if md.strip():
            yield md, None, table_idx

def chunk_text_stream(text_stream, sentences_per_chunk=5, overlap=2):
    buffer = ""
    sentences = []  # list of (sentence_text, page_num)
    current_page = None

    for text_part, page_num in text_stream:
        if page_num is not None:
            current_page = page_num

        buffer += " " + text_part
        new_sentences = sent_tokenize(buffer)

        if len(new_sentences) > 1:
            sentences.extend([(s, current_page) for s in new_sentences[:-1]])
            buffer = new_sentences[-1]

        while len(sentences) >= sentences_per_chunk:
            chunk_sents = sentences[:sentences_per_chunk]
            chunk = " ".join(s for s, _ in chunk_sents)
            page = chunk_sents[0][1]
            yield chunk.strip(), page
            sentences = sentences[sentences_per_chunk - overlap:]

    if buffer.strip():
        sentences.extend([(s, current_page) for s in sent_tokenize(buffer)])

    while len(sentences) > 0:
        chunk_sents = sentences[:sentences_per_chunk]
        chunk = " ".join(s for s, _ in chunk_sents)
        page = chunk_sents[0][1]
        yield chunk.strip(), page
        if len(sentences) <= sentences_per_chunk:
            break
        sentences = sentences[sentences_per_chunk - overlap:]

def yield_file_chunks(path, sentences_per_chunk=3, overlap=1):
    """Yield (text, page, metadata_dict) tuples for all content types.
    metadata_dict contains at minimum {"type": "text"|"table"}.
    """
    ext = path.lower()

    # ── Text chunks ──
    if ext.endswith(".txt"):
        stream = stream_txt(path)
    elif ext.endswith(".pdf"):
        stream = stream_pdf(path)
    elif ext.endswith(".docx"):
        stream = stream_docx(path)
    else:
        return

    for chunk, page in chunk_text_stream(stream, sentences_per_chunk, overlap):
        yield chunk, page, {"type": "text"}

    # ── Table chunks (PDF and DOCX only) ──
    if ext.endswith(".pdf"):
        for md_text, page_num, table_idx in extract_pdf_tables(path):
            yield md_text, page_num, {"type": "table", "table_index": table_idx}
    elif ext.endswith(".docx"):
        for md_text, page_num, table_idx in extract_docx_tables(path):
            yield md_text, page_num, {"type": "table", "table_index": table_idx}
